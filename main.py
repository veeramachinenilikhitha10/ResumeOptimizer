import os
from docx import Document
import pdfplumber

# ------- Utility Functions -------
def load_job_description(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_keywords(job_desc):
    common = set([
        'and','or','the','a','to','with','in','for','of','on',
        'is','are','we','looking','plus','as','an','at','our',
        'by','will','your','have','this','that','it','skills',
        'requirements','responsibilities','preferred','should'
    ])
    words = set(job_desc.lower().replace(',', '').replace('.', '').split())
    return list(words - common)

def parse_docx(file):
    doc = Document(file)
    all_text = "\n".join([para.text for para in doc.paragraphs])
    return doc, all_text

def parse_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def find_missing_keywords(resume_text, keywords):
    missing = []
    lower_resume = resume_text.lower()
    for kw in keywords:
        if kw not in lower_resume:
            missing.append(kw)
    return missing

def classify_keywords(missing_keywords):
    # Simple keyword classifier (extend as you like)
    skill_keywords = []
    proj_keywords = []
    exp_keywords = []
    base_skills = ['python','django','flask','html','css','javascript','aws','azure','google','sql','api','rest','docker','kubernetes','postgresql','mysql','cloud']
    base_actions = ['developed','implemented','collaborated','designed','built','deployed','tested','optimised']
    for kw in missing_keywords:
        if kw.lower() in base_skills:
            skill_keywords.append(kw)
        elif kw.lower() in base_actions:
            exp_keywords.append(kw)
        else:
            proj_keywords.append(kw)
    return skill_keywords, exp_keywords, proj_keywords

def integrate_keywords_docx(doc, skill_keywords, exp_keywords, proj_keywords):
    changes = []
    skill_added, exp_added, proj_added = False, False, False
    for para in doc.paragraphs:
        txt = para.text.lower()
        # Intelligent section mapping (works for most resumes)
        if ("skill" in txt or "technical" in txt) and skill_keywords and not skill_added:
            para.text += ", " + ", ".join(skill_keywords)
            changes.append(f"Added to Skills: {', '.join(skill_keywords)}")
            skill_added = True
        elif ("experience" in txt or "intern" in txt or "work" in txt) and exp_keywords and not exp_added:
            para.text += ". Keywords added: " + ", ".join(exp_keywords)
            changes.append(f"Added to Experience: {', '.join(exp_keywords)}")
            exp_added = True
        elif ("project" in txt) and proj_keywords and not proj_added:
            para.text += ". Keywords added: " + ", ".join(proj_keywords)
            changes.append(f"Added to Projects: {', '.join(proj_keywords)}")
            proj_added = True
    # Fallback: If nothing matched, append a keyword section at the end
    if not any([skill_added, exp_added, proj_added]):
        doc.add_paragraph("ATS Keywords Added: " + ", ".join(skill_keywords+exp_keywords+proj_keywords))
        changes.append("Added keywords at end (no matching sections found)")
    return doc, changes

def integrate_keywords_text_docx(text, skill_keywords, exp_keywords, proj_keywords):
    # For PDF resumes, generate new DOCX with keyword section at end
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("ATS Keywords Added: " + ", ".join(skill_keywords+exp_keywords+proj_keywords))
    changes = []
    if skill_keywords: changes.append(f"Added to Skills (plain text): {', '.join(skill_keywords)}")
    if exp_keywords: changes.append(f"Added to Experience (plain text): {', '.join(exp_keywords)}")
    if proj_keywords: changes.append(f"Added to Projects (plain text): {', '.join(proj_keywords)}")
    return doc, changes

def save_docx(doc, filename):
    doc.save(filename)

# ------- Main Flow -------
def main():
    print("=== AI Resume Optimiser (Advanced Version) ===")
    resume_file = input("Enter your resume file name (with extension, e.g., resume.docx or resume.pdf): ").strip()
    jd_file = input("Enter your job description file name (e.g., job_desc.txt): ").strip()

    if not os.path.exists(resume_file):
        print("Error: Resume file not found.")
        return
    if not os.path.exists(jd_file):
        print("Error: Job description file not found.")
        return

    job_desc = load_job_description(jd_file)
    keywords = extract_keywords(job_desc)

    if resume_file.lower().endswith(".docx"):
        doc, resume_text = parse_docx(resume_file)
        missing = find_missing_keywords(resume_text, keywords)
        skill_keywords, exp_keywords, proj_keywords = classify_keywords(missing)
        doc, changes = integrate_keywords_docx(doc, skill_keywords, exp_keywords, proj_keywords)
    elif resume_file.lower().endswith(".pdf"):
        resume_text = parse_pdf(resume_file)
        missing = find_missing_keywords(resume_text, keywords)
        skill_keywords, exp_keywords, proj_keywords = classify_keywords(missing)
        doc, changes = integrate_keywords_text_docx(resume_text, skill_keywords, exp_keywords, proj_keywords)
    else:
        print("Unsupported resume file type. Use .pdf or .docx")
        return

    save_docx(doc, "optimised_resume.docx")
    with open("change_log.txt", "w", encoding="utf-8") as logf:
        for change in changes:
            logf.write(change + "\n")

    print("\n--- Changes Applied ---")
    for change in changes:
        print(change)
    print("\nOptimised resume saved as 'optimised_resume.docx'")
    print("Change log saved as 'change_log.txt'")

if __name__ == "__main__":
    main()
